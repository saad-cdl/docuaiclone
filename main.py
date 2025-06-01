import os
import io
import requests
import base64
import logging
from flask import Flask, request, render_template, redirect, url_for, flash, jsonify, session, send_file
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import fitz 
import docx
from PIL import Image
from io import BytesIO
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from transformers import MBartTokenizer, TFMBartForConditionalGeneration, BlipProcessor, BlipForConditionalGeneration
import tensorflow as tf
import torch
import textract

# Set up logging
logging.basicConfig(level=logging.INFO)

# Flask app secret key for session management
app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Configure upload folder
app.config['UPLOAD_FOLDER'] = 'uploads'
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}

# Database setup
from db_setup import db, User, init_app  # Import db, User, and init_app from db_setup

# Initialize the database and login manager
init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'  # Redirects to the 'login' route if not logged in

# Check GPU availability
print("Num GPUs Available (TensorFlow): ", len(tf.config.list_physical_devices('GPU')))
print("Num GPUs Available (PyTorch): ", torch.cuda.device_count())

# Limit GPU memory usage in TensorFlow
gpus = tf.config.list_physical_devices('GPU')
if gpus:
    try:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)  # Allow memory growth
            # tf.config.experimental.set_virtual_device_configuration(
            #     gpu,
            #     [tf.config.experimental.VirtualDeviceConfiguration(memory_limit=8192)])  # Limit GPU memory to 8GB
    except RuntimeError as e:
        print(e)

# Set up TensorFlow for GPU usage
strategy = tf.distribute.MirroredStrategy()
print("Number of TensorFlow devices: {}".format(strategy.num_replicas_in_sync))

with strategy.scope():
    # Initialize MBart model and tokenizer
    model_name = "facebook/mbart-large-50"
    tokenizer = MBartTokenizer.from_pretrained(model_name)
    model = TFMBartForConditionalGeneration.from_pretrained(model_name)

# Initialize BLIP model and processor using PyTorch
blip_model_name = "Salesforce/blip-image-captioning-base"
blip_processor = BlipProcessor.from_pretrained(blip_model_name)
blip_model = BlipForConditionalGeneration.from_pretrained(blip_model_name)

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def encode_image_to_base64(image_bytes):
    return base64.b64encode(image_bytes).decode('utf-8')

def extract_text_and_images_from_pdf(pdf_path):
    text = ""
    images = []
    try:
        doc = fitz.open(pdf_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()

            # Extract images
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                images.append(image_bytes)
    except Exception as e:
        logging.error(f"Error extracting text and images from PDF: {e}")
    return text, images

def extract_text_and_images_from_docx(docx_path):
    text = ""
    images = []
    try:
        doc = docx.Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                images.append(image_data)
    except Exception as e:
        logging.error(f"Error extracting text and images from DOCX: {e}")
    return text, images

def extract_images_and_text(file):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
    file.save(file_path)

    if file.filename.endswith('.pdf'):
        return extract_text_and_images_from_pdf(file_path)
    elif file.filename.endswith('.docx'):
        return extract_text_and_images_from_docx(file_path)

def summarize_text(text):
    # You can force the operation on CPU to save GPU memory
    with tf.device('/CPU:0'):
        inputs = tokenizer(text, return_tensors="tf", max_length=1024, truncation=True)
        summary_ids = model.generate(inputs["input_ids"], max_length=250, min_length=100, length_penalty=2.0, num_beams=4, early_stopping=True)
        summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

def describe_image(image):
    # Offload most operations to CPU to save GPU memory
    image = Image.open(io.BytesIO(image))
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    blip_model.to(device)
    inputs = blip_processor(images=image, return_tensors="pt").to('cpu')  # Use CPU for processing
    out = blip_model.generate(**inputs.to(device))
    description = blip_processor.decode(out[0], skip_special_tokens=True)
    return description

def clear_upload_folder():
    upload_folder = app.config['UPLOAD_FOLDER']
    for filename in os.listdir(upload_folder):
        file_path = os.path.join(upload_folder, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)

# Function to create DOCX
def create_docx(summary, image_descriptions):
    from docx import Document
    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    if image_descriptions:
        doc.add_heading('Image Descriptions', level=1)
        for img_base64, description in image_descriptions:
            doc.add_paragraph(description)

    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'summary.docx')
    doc.save(output_path)
    return output_path

# Function to create PDF
def create_pdf(summary, image_descriptions):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'summary.pdf')
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica", 12)
    c.drawString(100, height - 50, "Summary")
    c.setFont("Helvetica", 10)
    c.drawString(100, height - 70, summary)

    if image_descriptions:
        c.setFont("Helvetica", 12)
        c.drawString(100, height - 150, "Image Descriptions")
        y = height - 170
        for img_base64, description in image_descriptions:
            c.setFont("Helvetica", 10)
            c.drawString(100, y, description)
            y -= 20

    c.save()
    return output_path

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/download')
@login_required
def download_file():
    filetype = request.args.get('filetype')
    if filetype not in ['docx', 'pdf']:
        logging.error("Invalid filetype")
        return "Invalid filetype.", 400

    filename = f'summary.{filetype}'
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    logging.info(f"Requested file path: {file_path}")

    if not os.path.isfile(file_path):
        logging.error(f"{filetype.upper()} file not found at {file_path}")
        return f"{filetype.upper()} file not found. Please try processing your file again.", 400

    try:
        logging.info(f"Sending file: {file_path}")
        response = send_file(file_path, as_attachment=True)
        os.remove(file_path)  # Remove the file after sending
        return response
    except Exception as e:
        logging.error(f"Error sending file: {e}")
        return "Error sending file.", 500

@app.route('/stop_execution', methods=['POST'])
def stop_execution():
    # Implement logic to stop any ongoing processing
    clear_upload_folder()
    return jsonify({'status': 'stopped'})

@app.route('/uploader')
@login_required
def upload_form():
    if not current_user.is_authenticated:
        return redirect(url_for('login'))

    clear_upload_folder()
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        text, images = extract_images_and_text(file)
        summary = summarize_text(text)

        image_descriptions = []
        for image in images:
            description = describe_image(image)
            image_descriptions.append((encode_image_to_base64(image), description))

        output_format = request.form.get('output_format', 'pdf')
        if output_format == 'pdf':
            output_path = create_pdf(summary, image_descriptions)
        elif output_format == 'docx':
            output_path = create_docx(summary, image_descriptions)
        else:
            return jsonify({'error': 'Invalid output format'}), 400

        return jsonify({'file_url': url_for('download_file', filetype=output_format)})

    return jsonify({'error': 'File not allowed'}), 400

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()

        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('home'))
        else:
            flash('Invalid email or password')
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        hashed_password = generate_password_hash(password, method='sha256')
        user = User(email=email, password=hashed_password)
        db.session.add(user)
        db.session.commit()
        flash('User registered successfully')
        return redirect(url_for('login'))

    return render_template('register.html')

if __name__ == "__main__":
    app.run(port=5000, debug=True)  # Run the Flask app on port 5000 with debug mode enabled
